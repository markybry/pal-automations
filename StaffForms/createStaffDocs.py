"""
Generate staff documents from Word templates for each house, organised by month.

Templates folder structure:
  ./Templates/01 Monthly/   monthly documents — one file per house per month
  ./Templates/02 Weekly/    weekly documents  — one file per house per Monday

Output folder: ./output/<house>/<MM Month YYYY>/01 Monthly/
                              or               /02 Weekly/

Placeholder replacement:
  {{house}}   - house name (e.g. "19 Bransley")
  {{month}}   - month name (e.g. "April") — monthly templates
  {{date}}    - DD/MM/YYYY in document body, DD-MM-YYYY in filename — weekly
  {{weekNum}} - week number within the month (1–4) — weekly

House filtering:
  Filenames starting with "{{house}}" are generated for every house.
  Filenames starting with a specific house name are generated for that house only.

Manifest:
  A .manifest.json file is stored in each output subfolder. It records a hash
  of each file at the time it was generated. On subsequent runs, files whose
  hash has changed since generation are assumed to have been manually edited
  and are skipped.

  .xlsx templates are skipped (no placeholder replacement needed).

Usage:
    python createStaffDocs.py [--start DD/MM/YYYY] [--months N]

    --start   Any date in the first month to generate (defaults to auto-detect
              from output folder, or current month if output is empty).
    --months  Number of months to generate (default: 1).
"""

import calendar
import hashlib
import json
import os
import shutil
import sys
import argparse
from datetime import datetime, timedelta

try:
    from docx import Document
except ImportError:
    print("python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

HOUSES = ["19 Bransley", "17 Bransley"]
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DIR = os.path.join(SCRIPT_DIR, "output")
TEMPLATES_DIR = os.path.join(SCRIPT_DIR, "Templates")
MANIFEST_FILE = ".manifest.json"


# ── date helpers ──────────────────────────────────────────────────────────────

def week_of_month(monday: datetime) -> int:
    return (monday.day - 1) // 7 + 1


def get_mondays_in_month(year: int, month: int) -> list:
    _, last_day = calendar.monthrange(year, month)
    return [
        datetime(year, month, day)
        for day in range(1, last_day + 1)
        if datetime(year, month, day).weekday() == 0
    ]


def next_month(year: int, month: int) -> tuple:
    return (year + 1, 1) if month == 12 else (year, month + 1)


def get_existing_month_folders(house_dir: str) -> set:
    existing = set()
    if os.path.isdir(house_dir):
        for name in os.listdir(house_dir):
            try:
                dt = datetime.strptime(name, "%m %B %Y")
                existing.add((dt.year, dt.month))
            except ValueError:
                try:
                    dt = datetime.strptime(name, "%B %Y")
                    existing.add((dt.year, dt.month))
                except ValueError:
                    pass
    return existing


# ── template helpers ──────────────────────────────────────────────────────────

def applicable_houses(template_name: str) -> list:
    """Return which houses this template should be generated for."""
    if template_name.startswith("{{house}}"):
        return HOUSES
    for house in HOUSES:
        if template_name.startswith(house):
            return [house]
    return HOUSES


def list_templates(subdir: str) -> list:
    """Return .docx and .xlsx templates in a Templates subfolder."""
    path = os.path.join(TEMPLATES_DIR, subdir)
    if not os.path.isdir(path):
        print(f"Warning: template folder not found: {path}")
        return []
    return [
        name for name in sorted(os.listdir(path))
        if os.path.splitext(name)[1].lower() in (".docx", ".xlsx")
        and not name.startswith(".")
    ]


# ── manifest ──────────────────────────────────────────────────────────────────

def file_hash(path: str) -> str:
    return hashlib.sha256(open(path, "rb").read()).hexdigest()


def load_manifest(folder: str) -> dict:
    path = os.path.join(folder, MANIFEST_FILE)
    if os.path.exists(path):
        with open(path) as f:
            return json.load(f)
    return {}


def save_manifest(folder: str, manifest: dict) -> None:
    with open(os.path.join(folder, MANIFEST_FILE), "w") as f:
        json.dump(manifest, f, indent=2)


def is_manually_modified(output_path: str, manifest: dict) -> bool:
    """Return True if the file exists and its hash differs from when it was generated."""
    if not os.path.exists(output_path):
        return False
    stored = manifest.get(os.path.basename(output_path))
    if stored is None:
        return False  # no prior record — treat as safe to overwrite
    return file_hash(output_path) != stored


# ── document replacement ──────────────────────────────────────────────────────

def replace_in_paragraph(paragraph, replacements: dict) -> None:
    for run in paragraph.runs:
        for placeholder, value in replacements.items():
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, value)

    full_text = "".join(run.text for run in paragraph.runs)
    if any(ph in full_text for ph in replacements):
        new_text = full_text
        for placeholder, value in replacements.items():
            new_text = new_text.replace(placeholder, value)
        if paragraph.runs:
            paragraph.runs[0].text = new_text
            for run in paragraph.runs[1:]:
                run.text = ""


def _replace_headers_footers(doc: "Document", replacements: dict) -> None:
    for section in doc.sections:
        for part in (
            section.header,
            section.footer,
            section.even_page_header,
            section.even_page_footer,
            section.first_page_header,
            section.first_page_footer,
        ):
            if part is not None:
                for paragraph in part.paragraphs:
                    replace_in_paragraph(paragraph, replacements)
                for table in part.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                replace_in_paragraph(paragraph, replacements)


def replace_in_doc(doc: "Document", replacements: dict) -> None:
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, replacements)
    _replace_headers_footers(doc, replacements)


def replace_in_shiftplan_doc(
    doc: "Document", week_monday: datetime, base_replacements: dict
) -> None:
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, base_replacements)

    for i, table in enumerate(doc.tables):
        if i < 7:
            day_date = week_monday + timedelta(days=i)
            table_replacements = {
                **base_replacements,
                "{{date}}": day_date.strftime("%d/%m/%Y"),
            }
        else:
            table_replacements = base_replacements

        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, table_replacements)

    _replace_headers_footers(doc, base_replacements)


# ── main ──────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="Generate staff documents organised by month for each house."
    )
    parser.add_argument(
        "--start",
        type=str,
        default=None,
        help=(
            "Any date (DD/MM/YYYY) in the first month to generate. "
            "Defaults to auto-detecting the next missing month from the output folder, "
            "or the current month if the output folder is empty."
        ),
    )
    parser.add_argument(
        "--months",
        type=int,
        default=1,
        help="Number of months to generate (default: 1).",
    )
    args = parser.parse_args()

    monthly_templates = list_templates("01 Monthly")
    weekly_templates = list_templates("02 Weekly")

    if not monthly_templates and not weekly_templates:
        print(f"No .docx templates found in: {TEMPLATES_DIR}")
        sys.exit(1)

    forced_start = None
    if args.start:
        try:
            base = datetime.strptime(args.start, "%d/%m/%Y")
            forced_start = (base.year, base.month)
        except ValueError:
            print("Invalid date format. Use DD/MM/YYYY.")
            sys.exit(1)

    today = datetime.today()
    print(f"Houses: {HOUSES}")
    print()

    total_created = total_skipped = 0

    for house in HOUSES:
        house_dir = os.path.join(OUTPUT_DIR, house)
        os.makedirs(house_dir, exist_ok=True)

        if forced_start:
            start_year, start_month = forced_start
        else:
            existing = get_existing_month_folders(house_dir)
            if existing:
                start_year, start_month = next_month(*max(existing))
            else:
                start_year, start_month = today.year, today.month

        cur_year, cur_month = start_year, start_month
        for _ in range(args.months):
            month_label = datetime(cur_year, cur_month, 1).strftime("%m %B %Y")
            month_name = datetime(cur_year, cur_month, 1).strftime("%B")
            month_dir = os.path.join(house_dir, month_label)
            os.makedirs(month_dir, exist_ok=True)

            print(f"{house} — {month_label}")

            mondays = get_mondays_in_month(cur_year, cur_month)
            manifest = load_manifest(month_dir)

            # ── monthly templates ─────────────────────────────────────────
            for template_name in monthly_templates:
                if house not in applicable_houses(template_name):
                    continue

                replacements = {
                    "{{house}}": house,
                    "{{month}}": month_name,
                    "{{month}": month_name,  # fallback for missing closing }}
                    "{{weekNum}}": "",
                }
                output_name = template_name
                for ph, val in replacements.items():
                    output_name = output_name.replace(ph, val)

                output_path = os.path.join(month_dir, output_name)
                rel = os.path.relpath(output_path, OUTPUT_DIR)

                if is_manually_modified(output_path, manifest):
                    print(f"  Skipped (modified): {rel}")
                    total_skipped += 1
                    continue

                template_path = os.path.join(TEMPLATES_DIR, "01 Monthly", template_name)
                doc = Document(template_path)
                replace_in_doc(doc, replacements)
                doc.save(output_path)
                manifest[output_name] = file_hash(output_path)

                print(f"  Created: {rel}")
                total_created += 1

            # ── weekly templates ──────────────────────────────────────────
            for template_name in weekly_templates:
                if house not in applicable_houses(template_name):
                    continue

                is_shiftplan = "Weekly Shiftplan" in template_name

                for week_monday in mondays:
                    week_num = week_of_month(week_monday)
                    date_filename = week_monday.strftime("%d-%m-%Y")
                    date_disp = week_monday.strftime("%d/%m/%Y")

                    base_replacements = {
                        "{{house}}": house,
                        "{{date}}": date_disp,
                        "{{weekNum}}": str(week_num),
                        "{{month}}": month_name,
                        "{{month}": month_name,  # fallback for missing closing }}
                    }
                    name_replacements = {
                        "{{house}}": house,
                        "{{date}}": date_filename,
                        "{{weekNum}}": str(week_num),
                        "{{month}}": month_name,
                        "{{month}": month_name,  # fallback for missing closing }}
                    }

                    output_name = template_name
                    for ph, val in name_replacements.items():
                        output_name = output_name.replace(ph, val)

                    output_path = os.path.join(month_dir, output_name)
                    rel = os.path.relpath(output_path, OUTPUT_DIR)

                    if is_manually_modified(output_path, manifest):
                        print(f"  Skipped (modified): {rel}")
                        total_skipped += 1
                        continue

                    template_path = os.path.join(TEMPLATES_DIR, "02 Weekly", template_name)
                    if os.path.splitext(template_name)[1].lower() == ".xlsx":
                        shutil.copy2(template_path, output_path)
                    else:
                        doc = Document(template_path)
                        if is_shiftplan:
                            replace_in_shiftplan_doc(doc, week_monday, base_replacements)
                        else:
                            replace_in_doc(doc, base_replacements)
                        doc.save(output_path)
                    manifest[output_name] = file_hash(output_path)

                    print(f"  Created: {rel}")
                    total_created += 1

            save_manifest(month_dir, manifest)

            cur_year, cur_month = next_month(cur_year, cur_month)

    print(f"\nDone — {total_created} files created, {total_skipped} skipped (manually modified)")


if __name__ == "__main__":
    main()
