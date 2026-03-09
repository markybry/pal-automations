import os
from docx import Document
from docx.shared import Inches

def list_files(folder_path):
    """List files (not directories) in the given folder."""
    return [f for f in os.listdir(folder_path) if os.path.isfile(os.path.join(folder_path, f))]

def create_docx_with_file_list(folder_path, output_path):
    # Create a new Word document
    doc = Document()
    doc.add_heading('File List - Reading Tracker', 0)

    files = list_files(folder_path)

    # Add a table: 1 header row + len(files) rows, 2 columns
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Filename (no extension)'
    hdr_cells[1].text = 'Read?'

    # Fill rows with filenames
    for filename in files:
        name_without_ext = os.path.splitext(filename)[0]
        row_cells = table.add_row().cells
        row_cells[0].text = name_without_ext
        row_cells[1].text = '☐'  # Unicode checkbox (unchecked)

    # Save the document
    doc.save(output_path)
    print(f"Document saved to: {output_path}")

# Example usage
folder = r"F:\pal files\OneDrive\PAL - Team\Policies and Procedures"       # Change this
output_docx = r'./.data/policiesToRead.docx'  # Change this
create_docx_with_file_list(folder, output_docx)
